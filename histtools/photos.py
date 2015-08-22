# -*- coding: utf-8 -*-
"""
Created on Fri Mar 13 08:34:31 2015

History Photo Tools

Suite of tools for working with environmental photos

TODO:
Clean up photo output procedure
    - Thumbnail output is ugly
    - File names are ugly
    - Enable sql query
Minimize function parameters

@author: bbatt
"""


import arcpy
from PIL import Image
import os
import docx

from shared import subtypes, eligdict, styledict


def extract_photos_from_fc(infc, out_folder):
    if not os.path.exists(out_folder):
        os.mkdir(out_folder)
    fldBLOB = 'DATA'  
    fldAttName = 'ATT_NAME'
    fldGlobID = 'REL_GLOBALID'
    flds = [fldBLOB, fldAttName, fldGlobID]    
    #Attachment table
    tbl = os.path.join("{}__ATTACH".format(infc))
    with arcpy.da.SearchCursor(tbl, flds) as cursor:
       for row in cursor:
          binaryRep = row[0]
          GlobID = row[2]
          uniquename = arcpy.CreateUniqueName("{}.jpg".format(GlobID), 
                                                  out_folder)
          open(uniquename,'wb').write(binaryRep.tobytes())    
          
def create_thumbnails(photo_folder):    
    os.chdir(photo_folder)
    for photo in os.listdir(photo_folder):
        if os.path.splitext(photo)[1].upper() == ".JPG":
            fullpath = os.path.abspath(photo)
            fullpathdir = os.path.dirname(fullpath)
            thumbpath = os.path.join(fullpathdir, "thumbnails")
            if not os.path.exists(thumbpath):
                os.makedirs(thumbpath)
            img = Image.open(fullpath)
            img.thumbnail((390,260), Image.ANTIALIAS)
            img.save(os.path.join(thumbpath, 
                                      os.path.basename(fullpath)), "JPEG")                                 
    return thumbpath
                                      
def getdomaindescription(gdb, subtype, codedvalue):
    dmns = arcpy.da.ListDomains(gdb)
    for dmn in dmns:
        if dmn.name == subtype:
            vals = dmn.codedValues
            val = vals.get(codedvalue)
            return val

def create_resource_table(infc, photo_folder, spatref, idxfeat=None, sql=None):
    #Define table template
    thumbpath = create_thumbnails(photo_folder)    
    document = docx.Document(os.path.join(os.path.dirname(__file__), 
                                          "../templates/template.docx"))
    gdb = os.path.dirname(infc)
    flds = ["ResourceID", "PropName", "Address", 
            "SHAPE@X", "SHAPE@Y",  "StrucType", "BldType", 
            "StyleType", "Eligibility", "ConstYr", "Notes", "GlobalID"]
    with arcpy.da.SearchCursor(infc, flds, sql, spatref) as cursor:
        tbl = document.add_table(rows=1, cols=9)
        hdr_cells  = tbl.rows[0].cells
        hdr_cells[1].text = "Resource Number / Field Map Reference"
        hdr_cells[2].text = "Name"
        hdr_cells[3].text = "Address"
        hdr_cells[4].text = "UTM E"
        hdr_cells[5].text = "UTM N"
        hdr_cells[6].text = "Type/Style"
        hdr_cells[7].text = "NRHP Evaluation"
        hdr_cells[8].text = "Notes"
        for row in sorted(cursor):
            #Declare variables for readability
            resid = "{}".format(row[flds.index("ResourceID")])
            resname = "{}".format(row[flds.index("PropName")])
            address = "{}".format(row[flds.index("Address")])
            easting = row[flds.index("SHAPE@X")]
            northing = row[flds.index("SHAPE@Y")]
            structure = row[flds.index("StrucType")]
            bldg = row[flds.index("BldType")]
            stylerow = row[flds.index("StyleType")]
            nrhp = "{}".format(row[flds.index("Eligibility")])
            constyr = "{}".format(row[flds.index("ConstYr")])
            notes = "{}".format(row[flds.index("Notes")])
            glblid = "{}".format(row[flds.index("GlobalID")])
            #add feature count logic
            row_cells = tbl.add_row().cells
            spatjoin = spatial_join(idxfeat, (easting, northing), 
                                    "PageName", spatref)
            if not spatjoin == None:
                row_cells[1].text = "{} / {}".format(resid, spatjoin)
            else:
                row_cells[1].text = "{}".format(resid)
            #temp picture holder
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            counterbool = True
            piccounter = 0
            pic = os.path.join(thumbpath, "{}.jpg".format(glblid))
            if os.path.exists(pic):
                hgt = 1400000
                run.add_picture(pic, width = hgt * 1.5, height = hgt)
                while counterbool == True:
                    pic = os.path.join(thumbpath, 
                                       "{}{}.jpg".format(glblid, piccounter))
                    if os.path.exists(pic):
                        hgt = 1400000
                        wdth = hgt * 1.5
                        run.add_picture(pic, width=wdth, height=hgt)
                        piccounter += 1
                    else:
                        counterbool = False
            else:
                pass
            try:
                bldgsub = subtypes.get(structure)
                bldgstyle = getdomaindescription(gdb, subtypes.get(structure), bldg)
                styletype = styledict.get(stylerow)
                eligibility = eligdict.get(nrhp)                   
                row_cells[2].text = resname
                cntys = r"C:\GIS\GDOT_Datasets.gdb\COUNTY"
                county = spatial_join(cntys, (easting, northing), 
                                      "NAME", spatref)
                if not county == None:
                    row_cells[3].text = "{} ({})".format(address, county)
                else:
                    row_cells[3].text = address
                row_cells[4].text = "{}".format(int(easting))
                row_cells[5].text = "{}".format(int(northing))
                row_cells[6].text = ", ".join((bldgsub, bldgstyle, styletype))
                row_cells[7].text = eligibility
                if not constyr == "None":
                    row_cells[8].text = "; ".join((constyr, '{}'.format(notes)))
                else:
                    row_cells[8].text = "{}".format(notes)
            except ValueError as e:
                print e.message
                print row
        document.save(os.path.join(photo_folder, "!tblReport.docx"))
#    os.remove(thumbpath)
    return document

def create_list_from_table_report(doc):
    #Assumes one table in history report
    tbl = doc.tables[0]
    dblist = []
    for row in tbl.rows:
        _toadd = []
        for cell in row.cells:
            _toadd.append(cell.text)
        dblist.append(_toadd)
    return dblist

def spatial_join(idxfeat, pnt_tup, idxfld, spatref):
    """
    On the fly spatial join for a given coordinate pair
    params:
    pnt_tup - UTM coordinate pair
    idxfeat - map index or other containing feature
    fld - field to label table with
    """
    #Check for valid feature input
    if not arcpy.Exists(idxfeat):
        return None
    else:
        pnt = arcpy.Point(pnt_tup[0], pnt_tup[1])    
        pntgeom = arcpy.PointGeometry(pnt, spatref)
        flds = [idxfld, 'SHAPE@']
        #Spatial references should be the same
        with arcpy.da.SearchCursor(idxfeat, flds, "", spatref) as rows:
            for row in rows:
                if row[1].contains(pntgeom):
                    return row[0]

def spatial_join_new(infc, inflds, idxfeat, idxfld, spatref):
    """
    On the fly spatial join for a given coordinate pair
    params:
    pnt_tup - UTM coordinate pair
    idxfeat - map index or other containing feature
    fld - field to label table with
    """
    flds = ["ResourceID", "PropName", "Address", 
            "SHAPE@X", "SHAPE@Y",  "StrucType", "BldType", 
            "StyleType", "Eligibility", "ConstYr", "Notes", "GlobalID"]
    rows = arcpy.da.SearchCursor(infc, flds)
    for row in rows:
        pnt = arcpy.Point(rows[3], rows[4])
        pntgeom = arcpy.PointGeometry(pnt, spatref)
        
    if not arcpy.Exists(idxfeat):
        return None
    else:
        pnt = arcpy.Point(pnt_tup[0], pnt_tup[1])    
        pntgeom = arcpy.PointGeometry(pnt, spatref)
        flds = [idxfld, 'SHAPE@']
        featdict = {}
        #Spatial references should be the same
        with arcpy.da.SearchCursor(idxfeat, flds, "", spatref) as rows:
            for row in rows:
                if row[1].contains(pntgeom):
                    pass
                    

def main():
    pass

if __name__ == "__main__":
    pass
    